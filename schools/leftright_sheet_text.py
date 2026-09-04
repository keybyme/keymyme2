"""Text extraction + best-effort line classification for
LeftRightSheetUpload -- one page of an already-existing Left & Right
turn-by-turn sheet, uploaded as a photo, PDF, or DOCX.

Images are OCR'd via the same Tesseract setup as vault/route_sheet_ocr.py
(free/offline, see settings.TESSERACT_CMD via that module's import-time
configuration) -- including HEIC/HEIF, once SchoolsConfig.ready() has
registered pillow-heif's Pillow opener. PDF/DOCX are digital documents, so
their text is pulled out directly rather than OCR'd -- MUCH more
accurate, since there's no image degradation to fight -- except a PDF
page with no embedded text (i.e. a scanned page saved as an image inside
the PDF), which falls back to rendering that one page and OCR'ing it like
a photo.

The line classifier is a draft-only guess, not a reliable one: every line
lands on the Edit page as an ordinary, already-editable LeftRightRow (see
LeftRightRowSaveView / leftright_form.html's autosave editing), which is
what makes a wrong guess here harmless rather than a reliability
requirement on the classifier itself.
"""

import io
import re

import pymupdf
import pytesseract
from docx import Document
from PIL import Image

from vault.route_sheet_ocr import extract_text as _extract_image_text

from .models import LeftRightRow

# Below this many characters of embedded text, a PDF page is treated as
# scanned/image-only (no real text layer) and OCR'd instead of trusted --
# a handful of stray characters (e.g. a page number) shouldn't count as
# "this page has real text".
MIN_PDF_PAGE_TEXT_CHARS = 10


def validate_sheet_file(file_obj, extension):
    """Raises ValueError with a user-facing detail if `file_obj` isn't
    actually a valid file of the type `extension` claims -- called before
    extract_sheet_text so a corrupted/mislabeled PDF or DOCX is rejected
    outright rather than saved with no text pulled from it. Doesn't cover
    images -- LeftRightSheetUploadView validates those itself via
    forms.ImageField (Pillow's own open-and-verify). Doesn't raise for an
    extraction problem on an otherwise-valid file (e.g. Tesseract being
    unavailable for a PDF's scanned page) -- that's extract_sheet_text's
    concern, not this one's."""
    file_obj.seek(0)
    if extension == "pdf":
        try:
            with pymupdf.open(stream=file_obj.read(), filetype="pdf"):
                pass
        except Exception as exc:
            raise ValueError(f"not a valid PDF ({exc})") from exc
    elif extension == "docx":
        try:
            Document(file_obj)
        except Exception as exc:
            raise ValueError(f"not a valid DOCX ({exc})") from exc
    file_obj.seek(0)


def extract_sheet_text(file_obj, extension):
    """Extracts text from one uploaded Left & Right sheet file. `extension`
    is the lowercased extension without the dot -- callers are expected to
    have already checked it's in LeftRightSheetUpload's allowed list.
    Raises on a file that can't be opened/parsed at all (a corrupted PDF/
    DOCX, or an image Pillow can't read); callers decide how to surface
    that (see LeftRightSheetUploadView)."""
    file_obj.seek(0)
    if extension == "pdf":
        return _extract_pdf_text(file_obj)
    if extension == "docx":
        return _extract_docx_text(file_obj)
    return _extract_image_text(file_obj)  # every other allowed extension is an image format


def _extract_pdf_text(file_obj):
    text_parts = []
    with pymupdf.open(stream=file_obj.read(), filetype="pdf") as pdf:
        for page in pdf:
            page_text = page.get_text().strip()
            if len(page_text) < MIN_PDF_PAGE_TEXT_CHARS:
                # No real text layer on this page -- it's a scanned image
                # saved into the PDF, not a digitally-created page. Render
                # it and OCR the rendered image instead, same engine as a
                # plain photo upload.
                pixmap = page.get_pixmap(dpi=300)
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                page_text = pytesseract.image_to_string(image).strip()
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_docx_text(file_obj):
    document = Document(file_obj)
    lines = [paragraph.text for paragraph in document.paragraphs]
    # Route sheets sometimes lay the guide out in a table rather than
    # plain paragraphs -- read those too, one classifier-friendly line per
    # row (cells joined with a space) rather than per cell.
    for table in document.tables:
        for row in table.rows:
            lines.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


# A line that OPENS with a house number ("123 Main St, City, MD 20874") is
# almost always a standalone address on the sheet, not a turn instruction
# -- "Turn left onto Main St" never starts with digits -- so this is what
# tells the two apart, not the presence of a street-suffix word alone
# (turn instructions name streets too).
ADDRESS_START_RE = re.compile(r"^\d{1,6}\s+\S")
# A short, fully-uppercase line ("AM ROUTE", "SCHOOL STOP") reads as a
# section heading rather than an instruction or address.
HEADING_RE = re.compile(r"^[A-Z0-9 .,'&/-]+$")
HEADING_MAX_WORDS = 6


def classify_line(line):
    """Returns (row_type, text, address) for one extracted line of text,
    per LeftRightRow's fields -- LINK rows carry the address in `address`
    since LeftRightDetailView turns that into a GPS link, not `text`."""
    if ADDRESS_START_RE.match(line):
        return LeftRightRow.RowType.LINK, "STOP @", line
    if HEADING_RE.match(line) and len(line.split()) <= HEADING_MAX_WORDS:
        return LeftRightRow.RowType.BOLD, line, ""
    return LeftRightRow.RowType.NORMAL, line, ""


def parse_sheet_text(text):
    """Splits extracted text into non-blank lines and classifies each one.
    Returns a list of (row_type, text, address) tuples, in reading order."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return [classify_line(line) for line in lines]
